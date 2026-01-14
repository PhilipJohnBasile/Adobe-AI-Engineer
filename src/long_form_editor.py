#!/usr/bin/env python3
"""
Long-Form Editor Backend

A Google Docs-style editor backend for generating long articles, blog posts,
and books, capable of accepting inline commands for AI-powered content generation.

Features:
- Document management with versioning
- Inline command processing (/expand, /rewrite, /continue, etc.)
- Section-based content generation
- Export to multiple formats (DOCX, PDF, HTML, Markdown)
- Auto-save and collaboration support
"""

import asyncio
import json
import logging
import os
import re
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict, field
from enum import Enum
from pathlib import Path
import uuid

logger = logging.getLogger(__name__)

# Optional imports for export functionality
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class DocumentStatus(Enum):
    """Document status types"""
    DRAFT = "draft"
    IN_PROGRESS = "in_progress"
    REVIEW = "review"
    PUBLISHED = "published"
    ARCHIVED = "archived"


class SectionType(Enum):
    """Section types in a document"""
    TITLE = "title"
    HEADING = "heading"
    SUBHEADING = "subheading"
    PARAGRAPH = "paragraph"
    LIST = "list"
    QUOTE = "quote"
    CODE = "code"
    IMAGE = "image"
    DIVIDER = "divider"


class InlineCommand(Enum):
    """Inline commands supported by the editor"""
    EXPAND = "/expand"
    REWRITE = "/rewrite"
    CONTINUE = "/continue"
    SUMMARIZE = "/summarize"
    SHORTEN = "/shorten"
    IMPROVE = "/improve"
    TRANSLATE = "/translate"
    TONE = "/tone"
    SIMPLIFY = "/simplify"
    ELABORATE = "/elaborate"
    BULLETS = "/bullets"
    PARAGRAPH = "/paragraph"
    OUTLINE = "/outline"
    CONCLUDE = "/conclude"
    INTRO = "/intro"
    EXAMPLE = "/example"
    STATS = "/stats"
    QUOTE = "/quote"


@dataclass
class Section:
    """Represents a section in a document"""
    id: str
    type: str
    content: str
    level: int = 1
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.content.split())


@dataclass
class DocumentVersion:
    """Represents a version of a document"""
    version_id: str
    content: str
    sections: List[Dict]
    created_at: str
    created_by: str
    message: str


@dataclass
class Document:
    """Represents a long-form document"""
    id: str
    title: str
    sections: List[Section]
    status: str = DocumentStatus.DRAFT.value
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    owner_id: str = "default"
    collaborators: List[str] = field(default_factory=list)
    version_history: List[DocumentVersion] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    word_count: int = 0
    target_word_count: int = 0
    voice_profile_id: Optional[str] = None

    def __post_init__(self):
        self.word_count = sum(s.word_count for s in self.sections)

    def get_full_content(self) -> str:
        """Get full document content as string"""
        parts = []
        for section in self.sections:
            if section.type == SectionType.TITLE.value:
                parts.append(f"# {section.content}\n")
            elif section.type == SectionType.HEADING.value:
                parts.append(f"\n## {section.content}\n")
            elif section.type == SectionType.SUBHEADING.value:
                parts.append(f"\n### {section.content}\n")
            elif section.type == SectionType.LIST.value:
                parts.append(section.content + "\n")
            elif section.type == SectionType.QUOTE.value:
                parts.append(f"\n> {section.content}\n")
            elif section.type == SectionType.DIVIDER.value:
                parts.append("\n---\n")
            else:
                parts.append(f"\n{section.content}\n")
        return "\n".join(parts)


class DocumentStore:
    """Persistent storage for documents"""

    def __init__(self, storage_path: str = "data/documents"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.documents: Dict[str, Document] = {}
        self._load_documents()

    def _load_documents(self):
        """Load documents from disk"""
        for doc_file in self.storage_path.glob("*.json"):
            try:
                with open(doc_file, 'r') as f:
                    data = json.load(f)
                # Reconstruct sections
                sections = [Section(**s) for s in data.get("sections", [])]
                data["sections"] = sections
                # Reconstruct version history
                versions = [DocumentVersion(**v) for v in data.get("version_history", [])]
                data["version_history"] = versions
                doc = Document(**data)
                self.documents[doc.id] = doc
            except Exception as e:
                logger.error(f"Error loading document {doc_file}: {e}")

    def save_document(self, doc: Document):
        """Save document to disk"""
        doc_file = self.storage_path / f"{doc.id}.json"
        data = asdict(doc)
        with open(doc_file, 'w') as f:
            json.dump(data, f, indent=2)
        self.documents[doc.id] = doc

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID"""
        return self.documents.get(doc_id)

    def delete_document(self, doc_id: str) -> bool:
        """Delete document"""
        if doc_id in self.documents:
            doc_file = self.storage_path / f"{doc_id}.json"
            if doc_file.exists():
                doc_file.unlink()
            del self.documents[doc_id]
            return True
        return False

    def list_documents(self, owner_id: str = None) -> List[Document]:
        """List all documents, optionally filtered by owner"""
        docs = list(self.documents.values())
        if owner_id:
            docs = [d for d in docs if d.owner_id == owner_id or owner_id in d.collaborators]
        return sorted(docs, key=lambda d: d.updated_at, reverse=True)


class LongFormEditor:
    """
    Google Docs-style long-form content editor with AI capabilities.

    Features:
    - Create and manage long-form documents
    - Process inline commands for AI-powered editing
    - Generate outlines and expand sections
    - Export to multiple formats
    - Version history and collaboration
    """

    # Command patterns
    COMMAND_PATTERN = re.compile(r'/(\w+)(?:\s+(.*))?$', re.MULTILINE)

    def __init__(self, text_generator=None, voice_profile_id: str = None):
        """Initialize the editor"""
        self.store = DocumentStore()
        self.voice_profile_id = voice_profile_id

        # Import text generator
        if text_generator:
            self.text_generator = text_generator
        else:
            try:
                from src.text_generator import TextGenerationEngine
                self.text_generator = TextGenerationEngine()
            except ImportError:
                self.text_generator = None
                logger.warning("Text generator not available")

    async def create_document(
        self,
        title: str,
        template: str = None,
        owner_id: str = "default",
        target_word_count: int = 0,
        voice_profile_id: str = None
    ) -> Document:
        """
        Create a new document.

        Args:
            title: Document title
            template: Optional template to use (blog, article, guide, etc.)
            owner_id: Owner user ID
            target_word_count: Target word count for the document
            voice_profile_id: Voice profile for consistent tone

        Returns:
            Created Document
        """
        doc_id = str(uuid.uuid4())[:8]

        # Create title section
        sections = [
            Section(
                id=str(uuid.uuid4())[:8],
                type=SectionType.TITLE.value,
                content=title,
                level=1
            )
        ]

        # Apply template if specified
        if template:
            template_sections = await self._get_template_sections(template, title)
            sections.extend(template_sections)

        doc = Document(
            id=doc_id,
            title=title,
            sections=sections,
            owner_id=owner_id,
            target_word_count=target_word_count,
            voice_profile_id=voice_profile_id or self.voice_profile_id
        )

        self.store.save_document(doc)
        logger.info(f"Created document: {doc_id} - {title}")

        return doc

    async def _get_template_sections(self, template: str, title: str) -> List[Section]:
        """Get template sections for document type"""
        templates = {
            "blog": [
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Introduction", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Write your introduction here]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Main Points", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Expand on your main points]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Conclusion", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Wrap up your article]", level=2),
            ],
            "article": [
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Overview", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Article overview]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Background", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Background information]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Analysis", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Your analysis]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Implications", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[What this means]", level=2),
            ],
            "guide": [
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="What You'll Learn", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.LIST.value, content="- Point 1\n- Point 2\n- Point 3", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Prerequisites", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[What readers need before starting]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Step 1: Getting Started", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[First step instructions]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Step 2: Next Steps", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Continue the guide]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Summary", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Key takeaways]", level=2),
            ],
            "ebook_chapter": [
                Section(id=str(uuid.uuid4())[:8], type=SectionType.QUOTE.value, content="[Chapter opening quote]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Chapter introduction]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Section 1", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Section content]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Section 2", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.PARAGRAPH.value, content="[Section content]", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.HEADING.value, content="Key Takeaways", level=2),
                Section(id=str(uuid.uuid4())[:8], type=SectionType.LIST.value, content="- Takeaway 1\n- Takeaway 2\n- Takeaway 3", level=2),
            ]
        }

        return templates.get(template, [])

    async def process_command(
        self,
        doc_id: str,
        command: str,
        context: str = None,
        section_id: str = None
    ) -> Dict[str, Any]:
        """
        Process an inline command in the document.

        Supported commands:
        - /expand [topic] - Expand on the current section
        - /rewrite [style] - Rewrite the current section
        - /continue - Continue writing from where you left off
        - /summarize - Summarize the current content
        - /shorten - Make the content more concise
        - /improve - Improve clarity and flow
        - /translate [language] - Translate to another language
        - /tone [tone] - Change the tone
        - /simplify - Simplify the language
        - /elaborate - Add more details
        - /bullets - Convert to bullet points
        - /paragraph - Convert bullets to paragraphs
        - /outline [topic] - Generate an outline
        - /conclude - Write a conclusion
        - /intro - Write an introduction
        - /example - Add examples
        - /stats - Add relevant statistics
        - /quote - Add a relevant quote

        Args:
            doc_id: Document ID
            command: The command string (e.g., "/expand marketing strategies")
            context: Current context/content around the command
            section_id: Optional section ID to apply command to

        Returns:
            Dict with generated content and metadata
        """
        doc = self.store.get_document(doc_id)
        if not doc:
            return {"error": "Document not found", "success": False}

        # Parse command
        match = self.COMMAND_PATTERN.match(command)
        if not match:
            return {"error": "Invalid command format", "success": False}

        cmd_name = match.group(1).lower()
        cmd_args = match.group(2) or ""

        # Get surrounding context if not provided
        if not context and section_id:
            section = next((s for s in doc.sections if s.id == section_id), None)
            if section:
                context = section.content

        # Process command
        result = await self._execute_command(cmd_name, cmd_args, context, doc)

        if result.get("success") and result.get("content"):
            # Create new section with generated content
            new_section = Section(
                id=str(uuid.uuid4())[:8],
                type=result.get("section_type", SectionType.PARAGRAPH.value),
                content=result["content"],
                metadata={"generated_by": f"/{cmd_name}", "args": cmd_args}
            )

            # Insert after specified section or at end
            if section_id:
                idx = next((i for i, s in enumerate(doc.sections) if s.id == section_id), -1)
                if idx >= 0:
                    doc.sections.insert(idx + 1, new_section)
            else:
                doc.sections.append(new_section)

            # Update document
            doc.updated_at = datetime.now().isoformat()
            doc.word_count = sum(s.word_count for s in doc.sections)
            self.store.save_document(doc)

            result["section_id"] = new_section.id

        return result

    async def _execute_command(
        self,
        cmd_name: str,
        cmd_args: str,
        context: str,
        doc: Document
    ) -> Dict[str, Any]:
        """Execute a specific command"""

        if not self.text_generator or not self.text_generator.available:
            return {"error": "Text generator not available", "success": False}

        handlers = {
            "expand": self._cmd_expand,
            "rewrite": self._cmd_rewrite,
            "continue": self._cmd_continue,
            "summarize": self._cmd_summarize,
            "shorten": self._cmd_shorten,
            "improve": self._cmd_improve,
            "translate": self._cmd_translate,
            "tone": self._cmd_tone,
            "simplify": self._cmd_simplify,
            "elaborate": self._cmd_elaborate,
            "bullets": self._cmd_bullets,
            "paragraph": self._cmd_paragraph,
            "outline": self._cmd_outline,
            "conclude": self._cmd_conclude,
            "intro": self._cmd_intro,
            "example": self._cmd_example,
            "stats": self._cmd_stats,
            "quote": self._cmd_quote,
        }

        handler = handlers.get(cmd_name)
        if not handler:
            return {"error": f"Unknown command: {cmd_name}", "success": False}

        try:
            return await handler(cmd_args, context, doc)
        except Exception as e:
            logger.error(f"Command execution error: {e}")
            return {"error": str(e), "success": False}

    async def _cmd_expand(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Expand on the current topic"""
        topic = args or context[:100] if context else doc.title

        result = await self.text_generator.generate(
            prompt=f"Expand on this topic with more details, examples, and explanations:\n\n{topic}\n\nPrevious context: {context[:500] if context else 'Start of document'}",
            system_prompt="You are an expert writer. Expand on the given topic with insightful, detailed content. Write 2-3 paragraphs.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "expand",
            "tokens_used": result.tokens_used
        }

    async def _cmd_rewrite(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Rewrite the content"""
        style = args or "improved clarity and flow"

        if not context:
            return {"error": "No content to rewrite", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Rewrite this content with {style}:\n\n{context}",
            system_prompt="You are an expert editor. Rewrite the content while maintaining the core message. Only output the rewritten text.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "rewrite",
            "tokens_used": result.tokens_used
        }

    async def _cmd_continue(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Continue writing from where left off"""
        full_context = doc.get_full_content()[-2000:]  # Last 2000 chars for context

        result = await self.text_generator.generate(
            prompt=f"Continue writing naturally from where this text ends:\n\n{full_context}",
            system_prompt="You are an expert writer. Continue the narrative seamlessly, maintaining the same tone and style. Write 2-3 paragraphs.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "continue",
            "tokens_used": result.tokens_used
        }

    async def _cmd_summarize(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Summarize content"""
        content_to_summarize = context or doc.get_full_content()

        result = await self.text_generator.summarize(
            text=content_to_summarize,
            format=args if args in ["bullets", "paragraph", "tldr"] else "paragraph"
        )

        return {
            "success": True,
            "content": result,
            "section_type": SectionType.PARAGRAPH.value if args != "bullets" else SectionType.LIST.value,
            "command": "summarize"
        }

    async def _cmd_shorten(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Make content more concise"""
        if not context:
            return {"error": "No content to shorten", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Make this content more concise while keeping the key points:\n\n{context}",
            system_prompt="You are an expert editor. Reduce the word count by 30-50% while maintaining all essential information. Be direct and remove fluff.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "shorten",
            "tokens_used": result.tokens_used
        }

    async def _cmd_improve(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Improve clarity and flow"""
        if not context:
            return {"error": "No content to improve", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Improve the clarity, flow, and impact of this content:\n\n{context}",
            system_prompt="You are an expert editor. Improve the writing quality while maintaining the message. Fix awkward phrasing, improve transitions, and enhance readability.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "improve",
            "tokens_used": result.tokens_used
        }

    async def _cmd_translate(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Translate content"""
        language = args or "Spanish"

        if not context:
            return {"error": "No content to translate", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Translate this to {language}:\n\n{context}",
            system_prompt=f"You are a professional translator. Translate accurately to {language} while maintaining the tone and meaning.",
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "translate",
            "tokens_used": result.tokens_used
        }

    async def _cmd_tone(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Change the tone"""
        target_tone = args or "professional"

        if not context:
            return {"error": "No content to change tone", "success": False}

        result = await self.text_generator.change_tone(context, target_tone)

        return {
            "success": True,
            "content": result,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "tone"
        }

    async def _cmd_simplify(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Simplify language"""
        if not context:
            return {"error": "No content to simplify", "success": False}

        target_level = args or "middle_school"
        result = await self.text_generator.simplify(context, target_level)

        return {
            "success": True,
            "content": result,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "simplify"
        }

    async def _cmd_elaborate(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Add more details"""
        topic = args or context[:100] if context else doc.title

        result = await self.text_generator.generate(
            prompt=f"Elaborate on this with more details, explanations, and depth:\n\n{topic}\n\nContext: {context[:500] if context else ''}",
            system_prompt="You are an expert writer. Add depth and detail to the topic. Include specific examples, explanations, and insights.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "elaborate",
            "tokens_used": result.tokens_used
        }

    async def _cmd_bullets(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Convert to bullet points"""
        if not context:
            return {"error": "No content to convert", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Convert this to clear, concise bullet points:\n\n{context}",
            system_prompt="Convert the content to bullet points. Each bullet should be one clear point. Use - for bullets.",
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.LIST.value,
            "command": "bullets",
            "tokens_used": result.tokens_used
        }

    async def _cmd_paragraph(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Convert bullets to paragraphs"""
        if not context:
            return {"error": "No content to convert", "success": False}

        result = await self.text_generator.generate(
            prompt=f"Convert these bullet points to flowing paragraphs:\n\n{context}",
            system_prompt="Convert the bullet points to well-written paragraphs. Maintain all information but make it flow naturally.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "paragraph",
            "tokens_used": result.tokens_used
        }

    async def _cmd_outline(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Generate an outline"""
        topic = args or doc.title

        result = await self.text_generator.generate(
            prompt=f"Create a detailed outline for: {topic}",
            system_prompt="Create a comprehensive outline with main sections and sub-points. Use ## for main sections and - for sub-points.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.LIST.value,
            "command": "outline",
            "tokens_used": result.tokens_used
        }

    async def _cmd_conclude(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Write a conclusion"""
        full_content = doc.get_full_content()

        result = await self.text_generator.generate(
            prompt=f"Write a compelling conclusion for this article:\n\n{full_content[-3000:]}",
            system_prompt="Write a conclusion that summarizes key points, reinforces the main message, and ends with a call to action or thought-provoking statement.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "conclude",
            "tokens_used": result.tokens_used
        }

    async def _cmd_intro(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Write an introduction"""
        topic = args or doc.title

        result = await self.text_generator.generate(
            prompt=f"Write an engaging introduction for an article about: {topic}",
            system_prompt="Write a compelling introduction that hooks the reader, establishes the topic's importance, and previews what they'll learn.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "intro",
            "tokens_used": result.tokens_used
        }

    async def _cmd_example(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Add examples"""
        topic = args or context[:100] if context else doc.title

        result = await self.text_generator.generate(
            prompt=f"Provide 2-3 specific, concrete examples for: {topic}",
            system_prompt="Provide real-world, specific examples that illustrate the concept. Make them relatable and practical.",
            voice_profile_id=doc.voice_profile_id
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "example",
            "tokens_used": result.tokens_used
        }

    async def _cmd_stats(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Add relevant statistics"""
        topic = args or context[:100] if context else doc.title

        result = await self.text_generator.generate(
            prompt=f"Provide relevant statistics and data points about: {topic}",
            system_prompt="Provide relevant statistics with context. Include percentages, numbers, and cite general sources where applicable. Note: These are illustrative and should be verified.",
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.PARAGRAPH.value,
            "command": "stats",
            "tokens_used": result.tokens_used
        }

    async def _cmd_quote(self, args: str, context: str, doc: Document) -> Dict[str, Any]:
        """Add a relevant quote"""
        topic = args or context[:100] if context else doc.title

        result = await self.text_generator.generate(
            prompt=f"Provide a relevant, inspiring quote about: {topic}",
            system_prompt="Provide a relevant quote from a notable figure. Include the attribution. If creating an original quote in the style of wisdom, note it as 'Inspired wisdom'.",
        )

        return {
            "success": True,
            "content": result.content,
            "section_type": SectionType.QUOTE.value,
            "command": "quote",
            "tokens_used": result.tokens_used
        }

    async def generate_outline(
        self,
        doc_id: str,
        topic: str = None,
        depth: int = 3
    ) -> List[Section]:
        """
        Generate a content outline for the document.

        Args:
            doc_id: Document ID
            topic: Topic to outline (defaults to document title)
            depth: Outline depth (1-3)

        Returns:
            List of generated Section objects
        """
        doc = self.store.get_document(doc_id)
        if not doc:
            return []

        topic = topic or doc.title

        prompt = f"""Create a detailed outline for: {topic}

Requirements:
- {5 + depth * 2} main sections
- Each main section has 2-3 sub-points
- Make it comprehensive and logical
- Use ## for main sections, ### for sub-sections, - for points"""

        result = await self.text_generator.generate(
            prompt=prompt,
            system_prompt="You are an expert content strategist. Create a comprehensive, well-structured outline.",
            voice_profile_id=doc.voice_profile_id
        )

        # Parse outline into sections
        sections = []
        lines = result.content.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('## '):
                sections.append(Section(
                    id=str(uuid.uuid4())[:8],
                    type=SectionType.HEADING.value,
                    content=line[3:],
                    level=2
                ))
            elif line.startswith('### '):
                sections.append(Section(
                    id=str(uuid.uuid4())[:8],
                    type=SectionType.SUBHEADING.value,
                    content=line[4:],
                    level=3
                ))
            elif line.startswith('- '):
                sections.append(Section(
                    id=str(uuid.uuid4())[:8],
                    type=SectionType.LIST.value,
                    content=line,
                    level=4
                ))

        # Add sections to document
        doc.sections.extend(sections)
        doc.updated_at = datetime.now().isoformat()
        self.store.save_document(doc)

        return sections

    async def expand_section(
        self,
        doc_id: str,
        section_id: str,
        target_words: int = 300
    ) -> Section:
        """
        Expand a section to target word count.

        Args:
            doc_id: Document ID
            section_id: Section ID to expand
            target_words: Target word count

        Returns:
            Expanded Section
        """
        doc = self.store.get_document(doc_id)
        if not doc:
            return None

        section = next((s for s in doc.sections if s.id == section_id), None)
        if not section:
            return None

        result = await self.text_generator.generate(
            prompt=f"Expand this into approximately {target_words} words:\n\n{section.content}",
            system_prompt="You are an expert writer. Expand the content with details, examples, and insights while maintaining coherence.",
            voice_profile_id=doc.voice_profile_id,
            max_tokens=target_words * 2
        )

        section.content = result.content
        section.word_count = len(result.content.split())
        section.updated_at = datetime.now().isoformat()

        self.store.save_document(doc)
        return section

    def create_version(
        self,
        doc_id: str,
        message: str = "Manual save",
        created_by: str = "user"
    ) -> DocumentVersion:
        """Create a version snapshot of the document"""
        doc = self.store.get_document(doc_id)
        if not doc:
            return None

        version = DocumentVersion(
            version_id=str(uuid.uuid4())[:8],
            content=doc.get_full_content(),
            sections=[asdict(s) for s in doc.sections],
            created_at=datetime.now().isoformat(),
            created_by=created_by,
            message=message
        )

        doc.version_history.append(version)
        self.store.save_document(doc)

        return version

    def restore_version(self, doc_id: str, version_id: str) -> bool:
        """Restore document to a previous version"""
        doc = self.store.get_document(doc_id)
        if not doc:
            return False

        version = next((v for v in doc.version_history if v.version_id == version_id), None)
        if not version:
            return False

        # Create backup of current state
        self.create_version(doc_id, f"Backup before restore to {version_id}", "system")

        # Restore sections
        doc.sections = [Section(**s) for s in version.sections]
        doc.updated_at = datetime.now().isoformat()
        self.store.save_document(doc)

        return True

    async def export(
        self,
        doc_id: str,
        format: str = "markdown"
    ) -> Tuple[bytes, str]:
        """
        Export document to various formats.

        Args:
            doc_id: Document ID
            format: Export format (markdown, html, docx, pdf, txt)

        Returns:
            Tuple of (file_bytes, filename)
        """
        doc = self.store.get_document(doc_id)
        if not doc:
            return None, None

        content = doc.get_full_content()
        safe_title = re.sub(r'[^\w\s-]', '', doc.title).strip().replace(' ', '_')

        if format == "markdown" or format == "md":
            return content.encode('utf-8'), f"{safe_title}.md"

        elif format == "txt":
            # Plain text without markdown formatting
            plain = re.sub(r'#+\s*', '', content)
            plain = re.sub(r'\*+', '', plain)
            plain = re.sub(r'>\s*', '', plain)
            return plain.encode('utf-8'), f"{safe_title}.txt"

        elif format == "html":
            if MARKDOWN_AVAILABLE:
                html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
                full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{doc.title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        blockquote {{ border-left: 3px solid #ccc; margin: 0; padding-left: 20px; color: #666; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
                return full_html.encode('utf-8'), f"{safe_title}.html"
            else:
                return content.encode('utf-8'), f"{safe_title}.html"

        elif format == "docx":
            if not DOCX_AVAILABLE:
                return None, "DOCX export not available. Install python-docx."

            docx_doc = DocxDocument()
            docx_doc.add_heading(doc.title, 0)

            for section in doc.sections:
                if section.type == SectionType.TITLE.value:
                    continue  # Already added as main heading
                elif section.type == SectionType.HEADING.value:
                    docx_doc.add_heading(section.content, level=1)
                elif section.type == SectionType.SUBHEADING.value:
                    docx_doc.add_heading(section.content, level=2)
                elif section.type == SectionType.QUOTE.value:
                    p = docx_doc.add_paragraph(section.content)
                    p.style = 'Quote'
                elif section.type == SectionType.LIST.value:
                    for item in section.content.split('\n'):
                        item = item.strip()
                        if item.startswith('- '):
                            docx_doc.add_paragraph(item[2:], style='List Bullet')
                else:
                    docx_doc.add_paragraph(section.content)

            # Save to bytes
            from io import BytesIO
            buffer = BytesIO()
            docx_doc.save(buffer)
            return buffer.getvalue(), f"{safe_title}.docx"

        elif format == "pdf":
            if not PDF_AVAILABLE:
                return None, "PDF export not available. Install reportlab."

            from io import BytesIO
            buffer = BytesIO()

            pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            for section in doc.sections:
                if section.type == SectionType.TITLE.value:
                    story.append(Paragraph(section.content, styles['Title']))
                elif section.type == SectionType.HEADING.value:
                    story.append(Spacer(1, 12))
                    story.append(Paragraph(section.content, styles['Heading1']))
                elif section.type == SectionType.SUBHEADING.value:
                    story.append(Spacer(1, 8))
                    story.append(Paragraph(section.content, styles['Heading2']))
                elif section.type == SectionType.QUOTE.value:
                    story.append(Paragraph(f"<i>{section.content}</i>", styles['Normal']))
                else:
                    story.append(Paragraph(section.content, styles['Normal']))
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            return buffer.getvalue(), f"{safe_title}.pdf"

        return None, None

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID"""
        return self.store.get_document(doc_id)

    def list_documents(self, owner_id: str = None) -> List[Dict[str, Any]]:
        """List all documents"""
        docs = self.store.list_documents(owner_id)
        return [
            {
                "id": d.id,
                "title": d.title,
                "status": d.status,
                "word_count": d.word_count,
                "updated_at": d.updated_at,
                "sections_count": len(d.sections)
            }
            for d in docs
        ]

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document"""
        return self.store.delete_document(doc_id)


# Demo function
async def demo_long_form_editor():
    """Demonstrate long-form editor capabilities"""
    print("Long-Form Editor Demo")
    print("=" * 50)

    editor = LongFormEditor()

    # Create a new document
    print("\n1. Creating new document...")
    doc = await editor.create_document(
        title="The Complete Guide to AI Content Creation",
        template="guide",
        target_word_count=2000
    )
    print(f"   Created: {doc.id} - {doc.title}")

    # List available commands
    print("\n2. Available inline commands:")
    for cmd in InlineCommand:
        print(f"   {cmd.value}")

    # Show document structure
    print(f"\n3. Document structure ({len(doc.sections)} sections):")
    for section in doc.sections[:5]:
        print(f"   - [{section.type}] {section.content[:50]}...")

    # Export capabilities
    print("\n4. Export formats available:")
    formats = ["markdown", "html", "txt"]
    if DOCX_AVAILABLE:
        formats.append("docx")
    if PDF_AVAILABLE:
        formats.append("pdf")
    print(f"   {', '.join(formats)}")

    print("\n Demo complete!")


if __name__ == "__main__":
    asyncio.run(demo_long_form_editor())
