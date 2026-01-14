#!/usr/bin/env python3
"""
Knowledge Base System

Upload and manage company documents, product specs, and style guides
so the AI creates factually accurate content specific to your business context.

Features:
- Document upload and parsing (PDF, DOCX, TXT, MD)
- Vector-based semantic search
- Content validation against knowledge base
- RAG (Retrieval Augmented Generation) support
"""

import asyncio
import json
import logging
import os
import re
import hashlib
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict, field
from pathlib import Path
import uuid

logger = logging.getLogger(__name__)

# Optional imports for advanced features
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentType:
    """Document type classifications"""
    PRODUCT_SPECS = "product_specs"
    BRAND_GUIDELINES = "brand_guidelines"
    COMPANY_INFO = "company_info"
    FAQ = "faq"
    CASE_STUDIES = "case_studies"
    COMPETITOR_INFO = "competitor_info"
    PRICING = "pricing"
    LEGAL = "legal"
    TECHNICAL_DOCS = "technical_docs"
    MARKETING_MATERIALS = "marketing_materials"


@dataclass
class KnowledgeChunk:
    """A chunk of knowledge from a document"""
    id: str
    document_id: str
    content: str
    chunk_index: int
    metadata: Dict[str, Any]
    embedding: List[float] = None

    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "document_id": self.document_id,
            "content": self.content,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata
        }


@dataclass
class KnowledgeDocument:
    """A document in the knowledge base"""
    id: str
    name: str
    doc_type: str
    file_path: str
    chunks: List[KnowledgeChunk]
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    owner_id: str = "default"


@dataclass
class SearchResult:
    """A search result from the knowledge base"""
    chunk: KnowledgeChunk
    score: float
    document_name: str
    document_type: str


class SimpleVectorStore:
    """Simple in-memory vector store for embeddings"""

    def __init__(self):
        self.embeddings: Dict[str, np.ndarray] = {}
        self.chunks: Dict[str, KnowledgeChunk] = {}

    def add(self, chunk_id: str, embedding: np.ndarray, chunk: KnowledgeChunk):
        """Add embedding to store"""
        self.embeddings[chunk_id] = embedding
        self.chunks[chunk_id] = chunk

    def search(self, query_embedding: np.ndarray, top_k: int = 5) -> List[Tuple[str, float]]:
        """Search for similar embeddings"""
        if not self.embeddings:
            return []

        similarities = []
        for chunk_id, embedding in self.embeddings.items():
            # Cosine similarity
            similarity = np.dot(query_embedding, embedding) / (
                np.linalg.norm(query_embedding) * np.linalg.norm(embedding)
            )
            similarities.append((chunk_id, float(similarity)))

        # Sort by similarity
        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:top_k]

    def remove(self, chunk_id: str):
        """Remove embedding from store"""
        if chunk_id in self.embeddings:
            del self.embeddings[chunk_id]
        if chunk_id in self.chunks:
            del self.chunks[chunk_id]

    def get_chunk(self, chunk_id: str) -> Optional[KnowledgeChunk]:
        """Get chunk by ID"""
        return self.chunks.get(chunk_id)


class KnowledgeBase:
    """
    Knowledge base for storing and querying company information.

    Features:
    - Document upload and parsing
    - Semantic search using embeddings
    - Content validation
    - RAG support for AI generation
    """

    # Chunk size for splitting documents
    CHUNK_SIZE = 500  # words
    CHUNK_OVERLAP = 50  # words

    def __init__(self, tenant_id: str = "default", storage_path: str = "data/knowledge_base"):
        self.tenant_id = tenant_id
        self.storage_path = Path(storage_path) / tenant_id
        self.storage_path.mkdir(parents=True, exist_ok=True)

        # Initialize embeddings model
        self.embedding_model = None
        if EMBEDDINGS_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("Loaded sentence transformer model")
            except Exception as e:
                logger.warning(f"Failed to load embedding model: {e}")

        # Initialize vector store
        self.vector_store = SimpleVectorStore()

        # Document storage
        self.documents: Dict[str, KnowledgeDocument] = {}
        self._load_documents()

    def _load_documents(self):
        """Load documents from disk"""
        index_file = self.storage_path / "index.json"
        if index_file.exists():
            try:
                with open(index_file, 'r') as f:
                    index_data = json.load(f)

                for doc_id, doc_data in index_data.items():
                    chunks = [KnowledgeChunk(**c) for c in doc_data.get("chunks", [])]
                    doc_data["chunks"] = chunks
                    doc = KnowledgeDocument(**doc_data)
                    self.documents[doc_id] = doc

                    # Rebuild vector store
                    for chunk in chunks:
                        if chunk.embedding and self.embedding_model:
                            self.vector_store.add(
                                chunk.id,
                                np.array(chunk.embedding),
                                chunk
                            )

                logger.info(f"Loaded {len(self.documents)} documents")
            except Exception as e:
                logger.error(f"Error loading knowledge base: {e}")

    def _save_index(self):
        """Save document index to disk"""
        index_file = self.storage_path / "index.json"
        index_data = {}

        for doc_id, doc in self.documents.items():
            doc_dict = asdict(doc)
            # Convert chunks
            doc_dict["chunks"] = [c.to_dict() for c in doc.chunks]
            index_data[doc_id] = doc_dict

        with open(index_file, 'w') as f:
            json.dump(index_data, f, indent=2)

    def _parse_document(self, file_path: Path) -> str:
        """Parse document content"""
        suffix = file_path.suffix.lower()

        try:
            if suffix == '.txt' or suffix == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                # Clean markdown
                if suffix == '.md':
                    content = re.sub(r'#+\s*', '', content)
                    content = re.sub(r'\*+', '', content)
                    content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
                return content

            elif suffix == '.pdf':
                if not PDF_AVAILABLE:
                    raise ImportError("PyPDF2 required for PDF parsing")
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text_parts = []
                    for page in reader.pages:
                        text_parts.append(page.extract_text())
                    return "\n\n".join(text_parts)

            elif suffix == '.docx':
                if not DOCX_AVAILABLE:
                    raise ImportError("python-docx required for DOCX parsing")
                doc = DocxDocument(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)

            elif suffix == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return json.dumps(data, indent=2)

            else:
                # Try as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            raise

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        words = text.split()
        chunks = []

        for i in range(0, len(words), self.CHUNK_SIZE - self.CHUNK_OVERLAP):
            chunk_words = words[i:i + self.CHUNK_SIZE]
            if chunk_words:
                chunks.append(" ".join(chunk_words))

        return chunks

    def _generate_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding for text"""
        if not self.embedding_model:
            return None

        try:
            embedding = self.embedding_model.encode(text)
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return None

    async def upload_document(
        self,
        file_path: str,
        doc_type: str,
        name: str = None,
        metadata: Dict[str, Any] = None,
        owner_id: str = "default"
    ) -> KnowledgeDocument:
        """
        Upload and index a document.

        Args:
            file_path: Path to document file
            doc_type: Document type classification
            name: Document name (defaults to filename)
            metadata: Additional metadata
            owner_id: Owner user ID

        Returns:
            Uploaded KnowledgeDocument
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Parse document content
        content = self._parse_document(path)

        if not content.strip():
            raise ValueError("Document has no content")

        # Generate document ID
        doc_id = hashlib.md5(f"{path.name}:{datetime.now().isoformat()}".encode()).hexdigest()[:12]

        # Chunk the content
        text_chunks = self._chunk_text(content)

        # Create chunk objects with embeddings
        chunks = []
        for i, chunk_text in enumerate(text_chunks):
            chunk_id = f"{doc_id}_{i}"
            embedding = self._generate_embedding(chunk_text)

            chunk = KnowledgeChunk(
                id=chunk_id,
                document_id=doc_id,
                content=chunk_text,
                chunk_index=i,
                metadata={"doc_type": doc_type, **(metadata or {})},
                embedding=embedding
            )
            chunks.append(chunk)

            # Add to vector store
            if embedding:
                self.vector_store.add(chunk_id, np.array(embedding), chunk)

        # Create document
        doc = KnowledgeDocument(
            id=doc_id,
            name=name or path.name,
            doc_type=doc_type,
            file_path=str(path.absolute()),
            chunks=chunks,
            metadata=metadata or {},
            owner_id=owner_id
        )

        # Store document
        self.documents[doc_id] = doc
        self._save_index()

        logger.info(f"Uploaded document: {doc_id} - {doc.name} ({len(chunks)} chunks)")
        return doc

    async def upload_text(
        self,
        content: str,
        doc_type: str,
        name: str,
        metadata: Dict[str, Any] = None,
        owner_id: str = "default"
    ) -> KnowledgeDocument:
        """
        Upload text content directly.

        Args:
            content: Text content
            doc_type: Document type classification
            name: Document name
            metadata: Additional metadata
            owner_id: Owner user ID

        Returns:
            Created KnowledgeDocument
        """
        # Save to temp file and upload
        temp_path = self.storage_path / f"temp_{uuid.uuid4().hex[:8]}.txt"
        with open(temp_path, 'w', encoding='utf-8') as f:
            f.write(content)

        try:
            doc = await self.upload_document(
                str(temp_path),
                doc_type,
                name,
                metadata,
                owner_id
            )
            return doc
        finally:
            # Clean up temp file
            if temp_path.exists():
                temp_path.unlink()

    async def query(
        self,
        question: str,
        doc_types: List[str] = None,
        top_k: int = 5
    ) -> List[SearchResult]:
        """
        Query the knowledge base.

        Args:
            question: Query question
            doc_types: Filter by document types
            top_k: Number of results to return

        Returns:
            List of SearchResult objects
        """
        if not self.embedding_model:
            # Fallback to keyword search
            return await self._keyword_search(question, doc_types, top_k)

        # Generate query embedding
        query_embedding = self.embedding_model.encode(question)

        # Search vector store
        results = self.vector_store.search(query_embedding, top_k * 2)

        # Build search results
        search_results = []
        for chunk_id, score in results:
            chunk = self.vector_store.get_chunk(chunk_id)
            if not chunk:
                continue

            # Get document info
            doc = self.documents.get(chunk.document_id)
            if not doc:
                continue

            # Filter by doc type if specified
            if doc_types and doc.doc_type not in doc_types:
                continue

            search_results.append(SearchResult(
                chunk=chunk,
                score=score,
                document_name=doc.name,
                document_type=doc.doc_type
            ))

        return search_results[:top_k]

    async def _keyword_search(
        self,
        question: str,
        doc_types: List[str] = None,
        top_k: int = 5
    ) -> List[SearchResult]:
        """Fallback keyword search when embeddings unavailable"""
        keywords = set(question.lower().split())
        results = []

        for doc in self.documents.values():
            if doc_types and doc.doc_type not in doc_types:
                continue

            for chunk in doc.chunks:
                chunk_words = set(chunk.content.lower().split())
                overlap = len(keywords & chunk_words)

                if overlap > 0:
                    score = overlap / len(keywords)
                    results.append(SearchResult(
                        chunk=chunk,
                        score=score,
                        document_name=doc.name,
                        document_type=doc.doc_type
                    ))

        # Sort by score
        results.sort(key=lambda x: x.score, reverse=True)
        return results[:top_k]

    async def validate_content(
        self,
        content: str,
        doc_types: List[str] = None
    ) -> Dict[str, Any]:
        """
        Validate generated content against knowledge base.

        Args:
            content: Content to validate
            doc_types: Document types to check against

        Returns:
            Validation result with accuracy score and issues
        """
        # Extract key claims from content
        sentences = re.split(r'[.!?]+', content)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

        validation_results = []

        for sentence in sentences[:10]:  # Limit to first 10 sentences
            # Query knowledge base for supporting info
            results = await self.query(sentence, doc_types, top_k=3)

            if results:
                best_match = results[0]
                validation_results.append({
                    "claim": sentence,
                    "supported": best_match.score > 0.5,
                    "confidence": best_match.score,
                    "source": best_match.document_name,
                    "context": best_match.chunk.content[:200]
                })
            else:
                validation_results.append({
                    "claim": sentence,
                    "supported": False,
                    "confidence": 0.0,
                    "source": None,
                    "context": None
                })

        # Calculate overall accuracy
        supported_count = sum(1 for r in validation_results if r["supported"])
        accuracy = supported_count / max(1, len(validation_results))

        # Identify issues
        issues = [r["claim"] for r in validation_results if not r["supported"]]

        return {
            "accuracy_score": round(accuracy, 2),
            "total_claims": len(validation_results),
            "supported_claims": supported_count,
            "unsupported_claims": len(issues),
            "validation_details": validation_results,
            "potential_issues": issues[:5]
        }

    def get_context_for_generation(
        self,
        topic: str,
        doc_types: List[str] = None,
        max_chunks: int = 5
    ) -> str:
        """
        Get relevant context for content generation (RAG).

        Args:
            topic: Topic to generate content about
            doc_types: Document types to include
            max_chunks: Maximum chunks to include

        Returns:
            Formatted context string for AI generation
        """
        # Run synchronously for simplicity
        loop = asyncio.new_event_loop()
        results = loop.run_until_complete(self.query(topic, doc_types, max_chunks))
        loop.close()

        if not results:
            return ""

        context_parts = []
        for result in results:
            context_parts.append(f"[Source: {result.document_name}]\n{result.chunk.content}")

        return "\n\n---\n\n".join(context_parts)

    def get_document(self, doc_id: str) -> Optional[KnowledgeDocument]:
        """Get document by ID"""
        return self.documents.get(doc_id)

    def list_documents(self, doc_type: str = None) -> List[Dict[str, Any]]:
        """List all documents"""
        docs = list(self.documents.values())
        if doc_type:
            docs = [d for d in docs if d.doc_type == doc_type]

        return [
            {
                "id": d.id,
                "name": d.name,
                "doc_type": d.doc_type,
                "chunks_count": len(d.chunks),
                "created_at": d.created_at
            }
            for d in docs
        ]

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document"""
        if doc_id not in self.documents:
            return False

        doc = self.documents[doc_id]

        # Remove chunks from vector store
        for chunk in doc.chunks:
            self.vector_store.remove(chunk.id)

        # Remove document
        del self.documents[doc_id]
        self._save_index()

        return True

    def get_document_types(self) -> List[Dict[str, Any]]:
        """Get available document types"""
        return [
            {"id": DocumentType.PRODUCT_SPECS, "name": "Product Specifications"},
            {"id": DocumentType.BRAND_GUIDELINES, "name": "Brand Guidelines"},
            {"id": DocumentType.COMPANY_INFO, "name": "Company Information"},
            {"id": DocumentType.FAQ, "name": "FAQs"},
            {"id": DocumentType.CASE_STUDIES, "name": "Case Studies"},
            {"id": DocumentType.COMPETITOR_INFO, "name": "Competitor Information"},
            {"id": DocumentType.PRICING, "name": "Pricing Information"},
            {"id": DocumentType.LEGAL, "name": "Legal Documents"},
            {"id": DocumentType.TECHNICAL_DOCS, "name": "Technical Documentation"},
            {"id": DocumentType.MARKETING_MATERIALS, "name": "Marketing Materials"},
        ]


# Demo function
async def demo_knowledge_base():
    """Demonstrate knowledge base capabilities"""
    print("Knowledge Base System Demo")
    print("=" * 50)

    kb = KnowledgeBase()

    # Show available document types
    print("\n1. Available document types:")
    for dt in kb.get_document_types():
        print(f"   - {dt['name']}")

    # Upload sample content
    print("\n2. Uploading sample content...")
    sample_content = """
    ContentAI Product Specifications

    ContentAI is an AI-powered content generation platform designed for marketing teams.

    Key Features:
    - GPT-4 powered text generation
    - 50+ pre-built templates
    - Brand voice learning
    - Multi-channel support (social, email, ads, blogs)
    - Team collaboration tools
    - API access for integrations

    Technical Specifications:
    - Supports 10+ languages
    - 99.9% uptime SLA
    - SOC 2 Type II compliant
    - GDPR and CCPA compliant
    - Rate limit: 1000 requests/hour (Pro plan)

    Pricing:
    - Starter: $29/month (100 generations)
    - Pro: $79/month (1000 generations)
    - Enterprise: Custom pricing
    """

    doc = await kb.upload_text(
        content=sample_content,
        doc_type=DocumentType.PRODUCT_SPECS,
        name="ContentAI Product Specs"
    )
    print(f"   Uploaded: {doc.id} - {doc.name} ({len(doc.chunks)} chunks)")

    # Query the knowledge base
    print("\n3. Querying knowledge base...")
    results = await kb.query("What are the key features of ContentAI?", top_k=3)
    print(f"   Found {len(results)} results:")
    for r in results[:2]:
        print(f"   - Score: {r.score:.2f} | {r.chunk.content[:100]}...")

    # Validate content
    print("\n4. Validating content...")
    test_content = "ContentAI supports 10 languages and offers GPT-4 powered generation. It costs $29/month for the starter plan."
    validation = await kb.validate_content(test_content)
    print(f"   Accuracy score: {validation['accuracy_score']}")
    print(f"   Supported claims: {validation['supported_claims']}/{validation['total_claims']}")

    # Get context for generation
    print("\n5. Getting context for generation...")
    context = kb.get_context_for_generation("ContentAI pricing")
    print(f"   Context length: {len(context)} chars")

    print("\n Demo complete!")


if __name__ == "__main__":
    asyncio.run(demo_knowledge_base())
